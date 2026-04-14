from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.5); s.bottom_margin = Inches(0.5)
    s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)

style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(10.5)
style.font.color.rgb = RGBColor(0x33,0x33,0x33)
style.paragraph_format.space_after = Pt(2); style.paragraph_format.space_before = Pt(0)

def heading(text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1a,0x1a,0x1a)
        r.font.size = Pt({0:24,1:22,2:13,3:11.5}.get(level,13))
    h.paragraph_format.space_before = Pt(6 if level>1 else 0)
    h.paragraph_format.space_after = Pt(3)
    return h

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold=True; r.font.size=Pt(10.5); r.font.name='Calibri'
        r2 = p.add_run(text); r2.font.size=Pt(10.5); r2.font.name='Calibri'
    else:
        r = p.add_run(text); r.font.size=Pt(10.5); r.font.name='Calibri'

def para(text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(size); r.font.name='Calibri'
    p.paragraph_format.space_after = Pt(2)

def line():
    p = doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    b = pBdr.makeelement(qn('w:bottom'), {qn('w:val'):'single',qn('w:sz'):'4',qn('w:space'):'1',qn('w:color'):'999999'})
    pBdr.append(b); pPr.append(pBdr)

def role(company, title, dates):
    p = doc.add_paragraph()
    r = p.add_run(company); r.bold=True; r.font.size=Pt(12); r.font.name='Calibri'
    p.paragraph_format.space_after = Pt(0)
    p2 = doc.add_paragraph()
    r2 = p2.add_run(title); r2.bold=True; r2.font.size=Pt(10.5); r2.font.name='Calibri'
    r3 = p2.add_run(f'  |  {dates}'); r3.font.size=Pt(10.5); r3.font.name='Calibri'; r3.font.color.rgb=RGBColor(0x66,0x66,0x66)
    p2.paragraph_format.space_after = Pt(3)

# HEADER
h = heading('Richard Williams', 0)
h.alignment = WD_ALIGN_PARAGRAPH.LEFT; h.paragraph_format.space_after = Pt(2)
c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = c.add_run('Seattle, WA  |  253.227.6494  |  richscottwill@gmail.com  |  LinkedIn')
r.font.size=Pt(10); r.font.name='Calibri'; r.font.color.rgb=RGBColor(0x66,0x66,0x66)
c.paragraph_format.space_after = Pt(4)
line()

# EXPERIENCE
heading('Professional Experience', 2)

# Amazon
role('Amazon Business', 'Marketing Manager, Paid Search', 'Sep 2022 \u2013 Present')
para('Own paid search strategy, forecasting, and experimentation across 10 global markets (US, CA, UK, DE, FR, IT, ES, JP, AU, MX). 40K+ monthly registrations, multi-million dollar annual spend, managed against weekly LTV targets (ie%CCP) and OP2 budget pacing.', italic=True, size=10)

bullet(' Designed rollout methodology and measurement framework. Launched 8/10 markets (7 at 100%). CA +18.5% regs vs plan; DE +16-20% lift. Playbook adopted as standard for all new market launches.', 'OCI (Offline Conversion Import):')
bullet(' Structured A/B testing program for ad copy, landing pages, bidding, and audiences across all markets. UK ad copy: +86% CTR, +31% regs. Every initiative follows hypothesis \u2192 design \u2192 execution \u2192 measurement \u2192 scale.', 'Experimentation:')
bullet(' Manage ie%CCP (incremental engagement-weighted cost per converted purchaser) weekly \u2014 connects acquisition to downstream purchasing and cohort quality. Built forecasting models for budget pacing and registration targets across 10 markets with weekly variance analysis.', 'LTV optimization:')
bullet(' Walmart Business grew Brand IS from 25% to 55% in US. Responded with bid caps + NB OCI efficiency, not budget escalation. Result: 32.9K regs (+16% vs plan, +68% YoY), CPA held.', 'Competitive response:')
bullet(' Took over 5 European markets during -15% YoY Q1 2024 headwinds. Improved to +0.1% by H2; other WW markets remained -23% to -35%. Levers: localized ad copy, OCI bidding, campaign consolidation.', 'EU5 turnaround:')
bullet(' Partnered with MarTech (LiveRamp/Adobe audience pipelines), Data Science (incrementality), Legal (F90 lifecycle targeting), MCS (landing pages), ABMA (engagement \u2014 13%\u219230% match rate).', 'Cross-functional:')
bullet(' Built automation for performance reporting, market monitoring, and cross-market analysis. Data pipelines (SQL, Python) reduced manual reporting and increased iteration velocity.', 'AI workflows:')
bullet(' Authored OP1 narrative (problem\u2192test\u2192result\u2192scale) across 5 workstreams. Developed Paid Search Testing Framework presented to Director-level leadership.', 'Strategic frameworks:')
bullet('+70% YoY conversions in 2023 via localized strategy and Yahoo JP expansion.', 'Japan: ')
bullet('Launched paid search for 5 partner teams \u2014 Enterprise, Engagement, Business Prime, Small Business Credit Card, Mobile App.', 'New programs: ')

# Getty Omnichannel
doc.add_paragraph()
role('Getty Images', 'Marketing Strategist, Omnichannel', 'Dec 2018 \u2013 Apr 2022')
para("Promoted into newly created role based on landing page initiative I led. Owned experimentation and performance across DG team's largest programs.", italic=True, size=10)

bullet(' Designed and ran omnichannel tests with Data Science, Ops, and Marketing. +20% above target on two largest order-volume programs.', 'Testing:')
bullet(' Developed paid channel strategies for subscription acquisition, connecting spend to retention and recurring revenue.', 'Subscriptions:')
bullet(' Owned $300K-$500K/year budget. Built SQL dashboard. Ran incrementality tests that established paid social as evergreen channel. Managed forecasting and pacing.', 'Paid Social:')
bullet(' Built global competitor campaign strategy. +6% immediate conversion lift. Adopted as standard across all markets.', 'Competitor framework:')
bullet(' DSA, SA360 data-driven attribution, brand holdout tests, display (GDN, DBM). Designed geo holdout methodology for incremental lift measurement.', 'Incrementality:')
bullet(' Standardized tagging across Paid Social platforms with Marketing Ops. Enabled cross-channel attribution.', 'Measurement infra:')

# Getty Specialist
p2 = doc.add_paragraph()
r2 = p2.add_run('Marketing Specialist, Paid Search & Paid Social'); r2.bold=True; r2.font.size=Pt(10.5); r2.font.name='Calibri'
r3 = p2.add_run('  |  Apr 2017 \u2013 Dec 2018'); r3.font.size=Pt(10.5); r3.font.name='Calibri'; r3.font.color.rgb=RGBColor(0x66,0x66,0x66)
p2.paragraph_format.space_after = Pt(3)

bullet(' Americas. +44% budget, +63% transactions, -12% CPA PoP.', '$12M/year paid search,')
bullet(' Developed tailored LP approach vs homepage-only. Results led directly to creation of Omnichannel role.', 'Landing page strategy:')
bullet(' Built business case for always-on social backed by incrementality testing. Changed team\u2019s channel strategy.', 'Paid Social advocacy:')

# Microsoft
doc.add_paragraph()
role('Microsoft', 'SEM Specialist, Paid Search', 'Jan 2015 \u2013 Apr 2017')
bullet(" Microsoft's AdWords, Bing Ads, Gemini. +82.3% YoY acquisitions, -51.4% YoY CPA.", 'In-house:')
bullet(' 35 Reseller Partner accounts ($6M/year). Oversight of $55M+/year portfolio.', 'Client-side:')

# Laplink
doc.add_paragraph()
role('Laplink Software', 'SEM Analyst', 'Apr 2013 \u2013 Jan 2015')
bullet('In-house PPC: +300% net revenue, +600% ROI YoY. International expansion (+15% net revenue). Weekly reporting to COO. Migrated email automation Act-On \u2192 Marketo.')

line()

# EDUCATION
heading('Education', 2)
p = doc.add_paragraph()
r1 = p.add_run('BA, Business Administration \u2014 Marketing'); r1.bold=True; r1.font.size=Pt(10.5); r1.font.name='Calibri'
r2 = p.add_run('  |  Seattle University, 2012'); r2.font.size=Pt(10.5); r2.font.name='Calibri'; r2.font.color.rgb=RGBColor(0x66,0x66,0x66)
line()

# SKILLS
heading('Technical Proficiencies', 2)
for label, val in [
    ('Paid Search: ', 'Google Ads, Bing Ads, Yahoo JP, SA360, OCI, AI Max'),
    ('Measurement: ', 'GA, Looker, Adobe Analytics, incrementality, geo holdouts, A/B testing, LTV/CAC, cohort analysis'),
    ('Data: ', 'SQL (Hive, Snowflake, DuckDB), Python, Excel, dashboards, forecasting'),
    ('AI/Automation: ', 'Agentic workflows, AI ad copy generation, automated reporting pipelines'),
    ('Platforms: ', 'LiveRamp, Adobe Audience Manager, Marketo, Meta/Instagram/LinkedIn Ads, GTM'),
    ('Global: ', '10 markets across NA, EU, JP, AU, LATAM \u2014 localized strategy, cross-regional coordination'),
]:
    p = doc.add_paragraph()
    r1 = p.add_run(label); r1.bold=True; r1.font.size=Pt(10); r1.font.name='Calibri'
    r2 = p.add_run(val); r2.font.size=Pt(10); r2.font.name='Calibri'
    p.paragraph_format.space_after = Pt(1)
line()

# RECOGNITION
heading('Recognition', 2)
for r in [
    'Highest performance review out of 70 (Getty, 2018)',
    'DG Team Year-End Award (Getty, 2019)',
    '3x Monthly Marketing Award (Getty, 2018\u20132021)',
    'Contract extensions during company-wide hiring freeze (Amazon, 2023)',
]:
    bullet(r)

doc.save('/tmp/Resume_RSW_OpenAI.docx')
print('Done')
