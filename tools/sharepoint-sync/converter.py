"""MarkdownConverter — convert markdown to SharePoint-compatible HTML or .docx bytes."""

from __future__ import annotations

import io
import re

import markdown
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models import ConversionError


class MarkdownConverter:
    """Converts markdown content to HTML (Site Mode) or .docx bytes (Directory Mode)."""

    # Markdown extensions for SharePoint-compatible HTML
    _MD_EXTENSIONS = ["tables", "fenced_code", "toc", "nl2br", "sane_lists"]

    # ------------------------------------------------------------------ #
    # HTML conversion (Site Mode)
    # ------------------------------------------------------------------ #

    def to_html(self, markdown_body: str) -> str:
        """Convert markdown to SharePoint-compatible HTML.

        Uses the ``markdown`` library with extensions for tables,
        fenced code blocks, and table-of-contents anchors.

        Raises ``ConversionError`` on failure.
        """
        try:
            html = markdown.markdown(
                markdown_body,
                extensions=self._MD_EXTENSIONS,
                output_format="html",
            )
            return html
        except Exception as exc:
            raise ConversionError(f"HTML conversion failed: {exc}") from exc

    # ------------------------------------------------------------------ #
    # DOCX conversion (Directory Mode)
    # ------------------------------------------------------------------ #

    def to_docx(self, markdown_body: str, metadata: dict) -> bytes:
        """Convert markdown to *.docx* bytes using ``python-docx``.

        Preserves headings, lists, tables, and code blocks.
        Includes the article title from *metadata* as the document title.

        Returns raw bytes of the .docx file.
        Raises ``ConversionError`` on failure.
        """
        try:
            doc = Document()

            # Title from metadata
            title = metadata.get("title", "Untitled")
            doc.add_heading(title, level=0)

            # Walk through lines and convert
            self._render_markdown_to_docx(doc, markdown_body)

            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except ConversionError:
            raise
        except Exception as exc:
            raise ConversionError(f"DOCX conversion failed: {exc}") from exc

    # ------------------------------------------------------------------ #
    # Internal helpers
    # ------------------------------------------------------------------ #

    def _render_markdown_to_docx(self, doc: Document, markdown_body: str) -> None:
        """Parse markdown line-by-line and add elements to the docx Document."""
        lines = markdown_body.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # Fenced code block
            if line.strip().startswith("```"):
                code_lines: list[str] = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                i += 1  # skip closing ```
                self._add_code_block(doc, "\n".join(code_lines))
                continue

            # Table (line contains | and next line is separator)
            if "|" in line and i + 1 < len(lines) and re.match(
                r"^\s*\|?[\s\-:|]+\|", lines[i + 1]
            ):
                table_lines = [line]
                i += 1
                while i < len(lines) and "|" in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                self._add_table(doc, table_lines)
                continue

            # Heading
            heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                doc.add_heading(text, level=min(level, 9))
                i += 1
                continue

            # Unordered list item
            ul_match = re.match(r"^(\s*)[-*+]\s+(.*)", line)
            if ul_match:
                text = ul_match.group(2)
                doc.add_paragraph(text, style="List Bullet")
                i += 1
                continue

            # Ordered list item
            ol_match = re.match(r"^(\s*)\d+\.\s+(.*)", line)
            if ol_match:
                text = ol_match.group(2)
                doc.add_paragraph(text, style="List Number")
                i += 1
                continue

            # Blank line — skip
            if not line.strip():
                i += 1
                continue

            # Regular paragraph
            doc.add_paragraph(line.strip())
            i += 1

    def _add_code_block(self, doc: Document, code: str) -> None:
        """Add a code block as a monospaced paragraph."""
        para = doc.add_paragraph()
        para.style = doc.styles["Normal"]
        run = para.add_run(code)
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    def _add_table(self, doc: Document, table_lines: list[str]) -> None:
        """Parse markdown table lines and add a docx Table."""
        def parse_row(line: str) -> list[str]:
            cells = line.strip().strip("|").split("|")
            return [c.strip() for c in cells]

        if len(table_lines) < 2:
            return

        header_cells = parse_row(table_lines[0])
        # Skip separator line (index 1)
        data_rows = [parse_row(l) for l in table_lines[2:] if l.strip()]

        num_cols = len(header_cells)
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = "Table Grid"

        # Header row
        for idx, cell_text in enumerate(header_cells):
            if idx < num_cols:
                table.rows[0].cells[idx].text = cell_text

        # Data rows
        for row_data in data_rows:
            row = table.add_row()
            for idx, cell_text in enumerate(row_data):
                if idx < num_cols:
                    row.cells[idx].text = cell_text
