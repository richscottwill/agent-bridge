#!/usr/bin/env python3
"""Validate all active state files against quality gates.

Checks: front-matter date, required sections, appendix count, null case,
data-through line, placeholder schema, source links, weasel words, docx integrity.

Usage:
    python3 validate_state_files.py              # validate all
    python3 validate_state_files.py --market MX  # validate only MX
    python3 validate_state_files.py --strict      # exit 1 on any warning

Requires: python-docx
"""

import argparse
import os
import re
import sys
from datetime import date

STATE_FILES_DIR = os.path.expanduser("~/shared/wiki/state-files")

REGISTRY = {
    "MX": "mx-paid-search-state",
    "AU": "au-paid-search-state",
    "WW": "ww-testing-state",
}

WEASEL_WORDS = [
    "should ", "might ", "could potentially", "a lot of", "many ",
    "warrants ", "primary suspect", "positive signal",
]

FALSE_POSITIVE_CONTEXTS = [
    "must not contain", "banned", "quality gates",
    "optimization should focus",  # quoting Brandon's directive
    "should produce",  # prescriptive, not hedging
    "should proceed",  # approved action
    "should not be credited",  # prescriptive negative
    "investigation should focus",  # prescriptive directive
    "many businesses close",  # factual holiday description
    "many mexican businesses",  # factual holiday description
]

REQUIRED_SECTIONS = [
    "introduction", "goals", "tenets",
    "state of the business", "lessons learned", "strategic priorities",
]


def validate_one(fname, today_str, strict=False):
    """Validate a single state file. Returns (pass_count, fail_count, warn_count)."""
    md_path = os.path.join(STATE_FILES_DIR, f"{fname}.md")
    docx_path = os.path.join(STATE_FILES_DIR, f"{fname}.docx")

    if not os.path.exists(md_path):
        print(f"  ❌ File not found: {md_path}")
        return 0, 1, 0

    with open(md_path) as f:
        content = f.read()

    content_lower = content.lower()
    passes, fails, warns = 0, 0, 0

    # 1. Front-matter date
    fm_match = re.search(r"updated:\s*(\S+)", content[:500])
    fm_date = fm_match.group(1) if fm_match else "NOT FOUND"
    if fm_date == today_str:
        print(f"  ✅ Date: {fm_date}")
        passes += 1
    else:
        print(f"  ⚠️  Date: {fm_date} (today is {today_str})")
        warns += 1

    # 2. Required sections
    missing = [s for s in REQUIRED_SECTIONS if s not in content_lower]
    if not missing:
        print(f"  ✅ Sections: all 6 present")
        passes += 1
    else:
        print(f"  ❌ Sections missing: {missing}")
        fails += 1

    # 3. Appendix count
    app_count = len(re.findall(r"^## Appendix", content, re.MULTILINE))
    if app_count >= 7:
        print(f"  ✅ Appendices: {app_count}")
        passes += 1
    else:
        print(f"  ⚠️  Appendices: {app_count} (expected ≥7)")
        warns += 1

    # 4. Null case
    if "cost of inaction" in content_lower:
        print(f"  ✅ Null case present")
        passes += 1
    else:
        print(f"  ❌ Null case MISSING")
        fails += 1

    # 5. Data-through line
    if "data through:" in content_lower or "data through:" in content:
        print(f"  ✅ Data-through line present")
        passes += 1
    else:
        print(f"  ❌ Data-through line MISSING")
        fails += 1

    # 6. Placeholder schema
    if "appendix h" in content_lower or "placeholder schema" in content_lower:
        print(f"  ✅ Placeholder schema (Appendix H)")
        passes += 1
    else:
        print(f"  ❌ Placeholder schema MISSING")
        fails += 1

    # 7. Source links
    if "appendix g" in content_lower and "source links" in content_lower:
        print(f"  ✅ Source links (Appendix G)")
        passes += 1
    else:
        print(f"  ❌ Source links MISSING")
        fails += 1

    # 8. Weasel words
    violations = []
    for w in WEASEL_WORDS:
        for m in re.finditer(re.escape(w), content, re.IGNORECASE):
            start = max(0, m.start() - 60)
            end = min(len(content), m.end() + 60)
            ctx = content[start:end].replace("\n", " ")
            is_fp = any(fp in ctx.lower() for fp in FALSE_POSITIVE_CONTEXTS)
            if not is_fp:
                violations.append((w.strip(), ctx.strip()))

    if not violations:
        print(f"  ✅ Weasel words: 0 violations")
        passes += 1
    else:
        print(f"  ⚠️  Weasel words: {len(violations)}")
        for w, ctx in violations[:3]:
            print(f"      \"{w}\" in: ...{ctx[:80]}...")
        warns += 1

    # 9. DOCX exists and is valid
    if os.path.exists(docx_path):
        try:
            from docx import Document
            doc = Document(docx_path)
            h_count = sum(1 for p in doc.paragraphs if p.style and "Heading" in p.style.name)
            t_count = len(doc.tables)
            size = os.path.getsize(docx_path)
            print(f"  ✅ DOCX: {size:,} bytes, {h_count} headings, {t_count} tables")
            passes += 1
        except Exception as e:
            print(f"  ❌ DOCX corrupt: {e}")
            fails += 1
    else:
        print(f"  ⚠️  DOCX not found (run convert_state_files.py)")
        warns += 1

    return passes, fails, warns


def main():
    parser = argparse.ArgumentParser(description="Validate state files against quality gates")
    parser.add_argument("--market", help="Validate only this market (MX, AU, WW)")
    parser.add_argument("--strict", action="store_true", help="Exit 1 on any warning")
    args = parser.parse_args()

    today_str = date.today().isoformat()
    print(f"STATE FILE VALIDATION — {today_str}")
    print("=" * 60)

    targets = REGISTRY
    if args.market:
        key = args.market.upper()
        if key not in REGISTRY:
            print(f"ERROR: Unknown market '{key}'")
            sys.exit(1)
        targets = {key: REGISTRY[key]}

    total_p, total_f, total_w = 0, 0, 0
    for market, fname in targets.items():
        print(f"\n{market}:")
        p, f, w = validate_one(fname, today_str, args.strict)
        total_p += p
        total_f += f
        total_w += w

    print(f"\n{'=' * 60}")
    print(f"TOTAL: {total_p} pass, {total_f} fail, {total_w} warn")

    if total_f > 0:
        print("❌ VALIDATION FAILED")
        sys.exit(1)
    elif total_w > 0 and args.strict:
        print("⚠️  WARNINGS in strict mode")
        sys.exit(1)
    else:
        print("✅ ALL CHECKS PASS")


if __name__ == "__main__":
    main()
